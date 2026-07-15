from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "SMA_Planning_Engine_Code_Study_Guide.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        set_cell_shading(hdr[idx], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 78, 121)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("SMA Planning Engine Code Study Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("A practical guide to the VB.NET WinForms project, its functions, SQL workflow, scheduling logic, and UI behavior.").italic = True
    doc.add_paragraph("Project path: C:\\Users\\AK\\Documents\\Codex\\2026-07-13\\sma-planning-engine-net-8")
    doc.add_paragraph("Audience: beginner to intermediate VB.NET developer who wants to understand what the code does and how each function connects to the screen.")

    doc.add_heading("1. How To Study This Code", level=1)
    add_numbered(doc, [
        "Start with Program.vb and LoginForm.vb to understand how the application opens.",
        "Read SMAPlannerForm.vb next because it is the first real business screen after login.",
        "Read SMASchedulerForm.vb in groups: project header, task buttons, task grid, task allocation, capacity planning, resource usage, save.",
        "Read SqlProjectRepository.vb alongside the forms because it explains where data comes from and where it is saved.",
        "Read ScheduleEngine.vb when you want to understand dates, dependencies, and 8-hour resource capacity.",
        "Read ScheduleTask.vb and TaskCatalogItem.vb last as the data shapes used by the grid and SQL task templates."
    ])

    doc.add_heading("2. Project Architecture", level=1)
    add_table(doc, ["Layer", "Main files", "Purpose"], [
        ("Startup and security", "Program.vb, LoginForm.vb", "Starts the app, checks username/password, then opens the planner form."),
        ("Main planner UI", "SMAPlannerForm.vb", "Searches SQL by Project ID, validates project metadata, then opens the scheduler."),
        ("Scheduler UI", "SMASchedulerForm.vb", "Shows tasks, allocation, capacity, resource usage, preview charts, save/refresh actions."),
        ("Scheduling logic", "ScheduleEngine.vb", "Recalculates start/finish dates, dependencies, working days, and 8-hour capacity."),
        ("SQL access", "SqlProjectRepository.vb", "Runs SQL SELECT/INSERT/UPDATE/DELETE and converts SQL rows into VB objects."),
        ("Catalog wrappers", "TaskCatalogService.vb, EmployeeCatalogService.vb, LiveProjectCatalogService.vb", "Small service classes that call the SQL repository."),
        ("Data models", "ScheduleTask.vb, TaskCatalogItem.vb, ProjectLibraryItem.vb, LiveProjectItem", "Objects used to move data between SQL, forms, grids, and charts."),
        ("Motion and visuals", "FormTransitionService.vb, SchedulerThemePalette, preview/chart classes", "Handles animated form transitions, themes, and pie chart previews.")
    ], widths=[1.4, 2.1, 3.0])

    doc.add_paragraph("Simple runtime flow:")
    add_code(doc, "Program.Main -> LoginForm -> SMAPlannerForm -> SQL Project ID lookup -> SMASchedulerForm -> SQL task templates -> user edits -> SaveProjectToSql -> SqlProjectRepository.SaveProject")

    doc.add_heading("3. Startup And Login", level=1)
    add_table(doc, ["Function", "Where", "What it does"], [
        ("Main", "Program.vb", "Enables WinForms visual styles, shows LoginForm, and opens SMAPlannerForm only when login succeeds."),
        ("LoginForm.New", "LoginForm.vb", "Builds the login screen, centers it, sets Enter key to the Login button, and fills username with current Windows user."),
        ("TryLogin", "LoginForm.vb", "Checks three rules: username is allowed, username matches the current Windows username, and password equals 12345."),
        ("CurrentWindowsUserName", "LoginForm.vb", "Reads WindowsIdentity.GetCurrent().Name, removes the domain prefix, and returns the local Windows username."),
        ("ShowLoginError", "LoginForm.vb", "Shows the error message, clears password, and keeps the user on the login form.")
    ], widths=[1.6, 1.4, 3.5])
    doc.add_paragraph("Allowed users are Sheik.Ahsan and Afzal.khan. Both currently use password 12345. Because the code also checks the Windows username, a user cannot simply type another allowed name unless the system username matches.")

    doc.add_heading("4. Planner Form Workflow", level=1)
    doc.add_paragraph("SMAPlannerForm is the first business screen. Its main job is to let the user type a Project ID and open the scheduler only when SQL has enough valid project data.")
    add_table(doc, ["Step", "Logic"], [
        ("1. User enters Project ID", "The Project ID is read from _liveProjectSearchBox.Text."),
        ("2. SQL connection check", "If _sqlRepository is Nothing, the app shows a SQL configuration warning."),
        ("3. Load project metadata", "GetProjectPlanningInfo(projectCode) reads Version_Table and Table_Project_Tracking."),
        ("4. Stop if missing", "If no project, missing project name, missing version, missing project size, invalid size, inactive project, or already planned, the scheduler does not open."),
        ("5. Build LiveProjectItem", "The SQL metadata is copied into a LiveProjectItem, including version, size, report type, pointcloud, tech pack, deed profile, shadow analysis, urgent small project."),
        ("6. Open scheduler", "SMASchedulerForm.LoadLiveProjectTemplate(projectToSchedule) loads SQL task templates and shows the scheduler with motion transition.")
    ], widths=[1.8, 4.7])

    doc.add_heading("Important SMAPlannerForm Functions", level=2)
    add_table(doc, ["Function", "Purpose"], [
        ("btnScheduleProject_Click", "Main Schedule Project button handler. Validates SQL data and opens scheduler."),
        ("RefreshSearchProjectSuggestions", "Updates search suggestions while typing."),
        ("SearchStoredProjects", "Finds already scheduled projects that match search text."),
        ("LoadProjectList", "Loads recent scheduled projects from SQL."),
        ("RefreshPlannerLists", "Reloads planner lists and summary counts."),
        ("ResolveSearchProject", "Chooses the typed or selected project suggestion."),
        ("BuildTaskReportFilter", "Builds report filter text such as BRE/ROL/Within, Shadow Analysis, Deed Profile."),
        ("BuildProjectDetailsText", "Builds the detail text shown for the selected SQL project."),
        ("OpenStoredProject", "Opens a previously saved project schedule from SQL."),
        ("SetPlannerStatus", "Updates the planner status label.")
    ], widths=[2.2, 4.3])

    doc.add_heading("5. SQL Project Lookup", level=1)
    doc.add_paragraph("Project lookup is intentionally SQL-only. The current code should not silently assume Version 1.0 or Project Size Small when SQL is missing values.")
    add_code(doc, "SELECT TOP 1 v.[Project Name], v.[Project ID at SMA], v.[Version], ... t.[Project Size]\nFROM dbo.Version_Table v\nLEFT JOIN dbo.Table_Project_Tracking t ON t.[Project ID at SMA] = v.[Project ID at SMA]\nWHERE v.[Project ID at SMA] = @ProjectIdAtSma")
    add_table(doc, ["SQL field", "VB property", "Why it matters"], [
        ("Version_Table.[Project Name]", "ProjectName", "Shown in the UI and saved with schedule rows."),
        ("Version_Table.[Version]", "VersionNumber", "Used to decide project type: whole number is New, .1 is Feedback."),
        ("Table_Project_Tracking.[Project Size]", "ProjectSize", "Controls which task hours are pulled: Small, Medium, Large, Very Large."),
        ("Report BRE / ROL / Within", "ReportType", "Controls which task template rows match the project report."),
        ("Is_Pointcloud, Teck Pack, Deed Profile, Shadow_Analysis, Urgent Small Projects", "Theme/detail flags", "Drive UI theme and project details display."),
        ("IsPlanned, Active", "Planning rules", "Prevents scheduling inactive projects or projects already planned.")
    ], widths=[2.4, 1.8, 2.3])

    doc.add_heading("6. Task Loading Logic", level=1)
    doc.add_paragraph("Tasks are loaded from SQL through TaskCatalogService and SqlProjectRepository. Runtime fallback task lists have been removed.")
    add_table(doc, ["Function", "Logic"], [
        ("TaskCatalogService.LoadTemplateTasks", "Calls _sqlRepository.LoadTaskTemplates and keeps only tasks where HoursForSize(projectSize) is greater than zero."),
        ("SqlProjectRepository.LoadTaskTemplates", "Resolves dbo.Task_Template and filters by project type, project size, active flag, and report type."),
        ("LoadTaskTemplatesFromDatabaseDesign", "Reads [Task ID], [Task Name], [Small (hrs)], [Medium (hrs)], [Large (hrs)], [Very Large (hrs)], [Project Type], [Task Order], [Type of Report]."),
        ("TaskCatalogItem.HoursForSize", "Returns the correct hours column for Small, Medium, Large, or Very Large. Unknown size returns 0."),
        ("SMASchedulerForm.LoadTemplateTasks", "Creates ScheduleTask rows from SQL catalog rows and totals project hours."),
        ("CreateTemplateScheduleTask", "Converts one SQL task template into one visible scheduler task row.")
    ], widths=[2.3, 4.2])
    doc.add_paragraph("If the SQL task template query returns no matching tasks, the scheduler throws a clear message: No tasks available for this project.")

    doc.add_heading("7. Version And Project Type Logic", level=1)
    add_table(doc, ["Version number", "Project type", "Meaning"], [
        ("1.0, 2.0, 3.0", "New", "Whole number version means new project task templates should be used."),
        ("1.1, 2.1, 3.1", "Feedback", "Decimal .1 version means feedback task templates should be used."),
        ("Missing version", "Stop", "The app should not open scheduler; user must fix SQL."),
        ("Unexpected version", "Default behavior in helper", "ProjectTypeFromVersion currently returns New if it cannot parse a known decimal pattern. This is a helper fallback, not a replacement for required SQL version validation.")
    ], widths=[1.5, 1.3, 3.7])

    doc.add_heading("8. Project Size And Hours Logic", level=1)
    doc.add_paragraph("Project size comes from SQL. It decides which hours column is used from Task_Template.")
    add_table(doc, ["Project size", "Task_Template hours column", "Used by"], [
        ("Small", "[Small (hrs)]", "TaskCatalogItem.SmallHours"),
        ("Medium", "[Medium (hrs)]", "TaskCatalogItem.MediumHours"),
        ("Large", "[Large (hrs)]", "TaskCatalogItem.LargeHours"),
        ("Very Large", "[Very Large (hrs)]", "TaskCatalogItem.VeryLargeHours")
    ], widths=[1.4, 2.2, 2.9])
    doc.add_paragraph("When a task is added from SQL, ResourceHours is set to the requested hours for that project size. DurationDays is calculated as hours divided by 8.")
    add_code(doc, "DurationFromHours(hours) = hours / 8")

    doc.add_heading("9. Scheduler Form Main Button Logic", level=1)
    add_table(doc, ["Button/function", "What happens"], [
        ("btnAddTask_Click", "Adds a blank manual task row with zero resource hours."),
        ("AddTaskFromCatalog", "Adds a selected SQL catalog task, applies project-size hours, and recalculates."),
        ("btnDelete_Click", "Deletes the selected task from the current in-memory schedule, removes dependencies pointing to it, renumbers tasks, and recalculates."),
        ("btnMoveUp_Click / btnMoveDown_Click", "Moves the selected task and remaps predecessors so links follow the new order."),
        ("btnLink_Click", "Links the selected task to the previous task using FS dependency."),
        ("btnUnlink_Click", "Clears predecessor dependency for selected task."),
        ("btnMilestone_Click", "Toggles Milestone text and sets duration to 1 working day."),
        ("btnSave_Click", "Validates there are tasks, then calls SaveProjectToSql."),
        ("btnRefreshCapacity_Click", "Commits grid edits, recalculates, saves SQL silently, refreshes workspace tabs, and shows result message.")
    ], widths=[2.2, 4.3])

    doc.add_heading("10. Scheduler Date, Dependency, And 8-Hour Logic", level=1)
    doc.add_paragraph("ScheduleEngine and SMASchedulerForm both apply working-day logic. The engine recalculates the whole task list; the form normalizes user edits in the grid.")
    add_table(doc, ["Function", "Logic"], [
        ("ScheduleEngine.Recalculate", "Cleans task values, normalizes working dates, applies dependencies, finds resource capacity, and updates StartDate/FinishDate."),
        ("ParseDependencies", "Reads predecessors like 3, 3FS, 3SS, 3FF, 3SF."),
        ("RequiredStartForDependency", "Calculates start date based on FS, SS, FF, or SF dependency."),
        ("FindStartWithCapacity", "Moves the start date forward until assigned resources have capacity."),
        ("AllocateTaskHours", "Spreads resource hours across working days using available hours."),
        ("AvailableHours", "Returns 8 minus already used hours for that employee/date. Weekends are zero unless Saturday option is enabled in the form."),
        ("NormalizeGridEdit", "When user edits start/finish/duration/resource fields, it recalculates dates and distribution."),
        ("WarnIfResourcesExceedDailyCapacity", "Warns/highlights when a resource is over daily capacity.")
    ], widths=[2.2, 4.3])
    doc.add_paragraph("Important restriction: the current scheduling capacity rule is 8 hours per resource per day. If one resource already has planned hours on a date, only the remaining hours are available.")

    doc.add_heading("11. Resource Assignment Logic", level=1)
    add_table(doc, ["Data field", "Example", "Meaning"], [
        ("AssignedTo", "Afzal.khan; Sheik.Ahsan", "The selected employee names shown in the main task grid."),
        ("ResourceNames", "Afzal.khan; Sheik.Ahsan", "Normalized resource list used by views and schedule calculations."),
        ("ResourceAllocations", "Afzal.khan=4; Sheik.Ahsan=4", "Total hours per employee for this task."),
        ("DailyResourceAllocations", "Afzal.khan|2026-07-15=4", "Exact planned hours by employee and date."),
        ("ResourceHours", "8", "Total planned hours for the task.")
    ], widths=[1.8, 2.1, 2.6])
    add_table(doc, ["Function", "Purpose"], [
        ("NormalizeResourceList", "Cleans employee names and removes duplicates."),
        ("KeepOnlySelectedResources", "When selected employees change, removes unselected resources from allocation strings."),
        ("SyncTaskResourceDistribution", "Splits task ResourceHours evenly between selected employees."),
        ("BuildEvenResourceAllocations", "Calculates equal hours per selected resource and keeps rounding sane."),
        ("BuildDefaultDailyAllocations", "Creates date-by-date allocation values for task/resource combinations."),
        ("SetTaskResourceHoursOnDate", "Updates one resource's hours for one date."),
        ("SetTaskTotalHoursOnDate", "Updates total task hours on one date and distributes across selected resources.")
    ], widths=[2.3, 4.2])

    doc.add_heading("12. Task Allocation, Task Usage, And Resource Usage", level=1)
    add_table(doc, ["View", "What it shows", "Main logic"], [
        ("Task Allocation", "Tasks and resource distribution/pie preview.", "BuildPreviewModel and BuildResourceContributionSlices summarize task hours and resource contribution."),
        ("Task Usage", "Task rows across dates.", "TaskHoursOnDate and DistributedTaskHoursOnDate calculate how many hours are planned on each date."),
        ("Resource Usage", "Resource rows across dates.", "PlannedHoursForResource and AvailableHoursForResource compare planned hours against available capacity."),
        ("Resource Utilization", "Usage summary and highlights.", "StyleResourceUtilizationCell and ApplyResourceUtilizationColor mark load levels."),
        ("Capacity Planning", "SQL-backed employees, available hours, planned hours, projects, and tasks by date.", "LoadCapacityPlanningData and RenderSqlCapacityPlanningGrid read Employee_Capacity and Project Schedule Table.")
    ], widths=[1.6, 2.3, 2.6])

    doc.add_heading("13. Pie Chart / Preview Logic", level=1)
    doc.add_paragraph("The pie chart panels are custom WinForms drawing panels. They do not come from SQL directly; they are built from the current task list after SQL tasks are loaded.")
    add_table(doc, ["Function/class", "What it does"], [
        ("BuildPreviewModel", "Builds the data model for chart display: slices, totals, task count, duration, assigned hours."),
        ("BuildResourceContributionSlices", "Groups resource allocations and calculates slice values."),
        ("ResourceAllocationsForPreview", "Parses a task's ResourceAllocations string for charting."),
        ("PlannerPieChartPanel.UpdateModel", "Receives the model and redraws the chart."),
        ("PlannerLegendGrid", "Creates the chart legend grid."),
        ("BuildPlannerLegendRows", "Turns chart slices into rows shown beside the chart.")
    ], widths=[2.3, 4.2])

    doc.add_heading("14. Capacity Planning SQL Logic", level=1)
    doc.add_paragraph("Capacity Planning can read directly from SQL for the selected date range. It shows active employees, availability, planned work, projects, and task rows.")
    add_table(doc, ["SQL source", "Used for"], [
        ("Employees_Master", "Employee ID, names, and Active=True filtering."),
        ("Employee_Capacity", "Available hours by employee/date."),
        ("Project Schedule Table", "Planned hours, task names, task order, project IDs, schedule dates."),
        ("Version_Table", "Project name and active/planned project metadata.")
    ], widths=[2.0, 4.5])
    add_table(doc, ["Function", "Purpose"], [
        ("UpdateCapacityPlanningGrid", "Builds the capacity grid columns based on selected date range and chooses SQL data if available."),
        ("CapacityFilterStartDate / CapacityFilterFinishDate", "Reads date picker values."),
        ("LoadSqlCapacityPlanningData", "Calls SQL repository for active employees, availability, and planned assignments."),
        ("RenderSqlCapacityPlanningGrid", "Applies selected employees and builds the display rows."),
        ("AddSqlCapacityAvailableRow", "Shows available hours by date for one employee."),
        ("AddSqlCapacityPlannedRow", "Shows total planned hours by date for one employee."),
        ("AddSqlCapacityProjectRow", "Shows project-level planned hours under the employee."),
        ("AddSqlCapacityTaskRow", "Shows task-level planned hours under the project."),
        ("StyleUsageSummaryCell", "Colors cells: over-capacity values are highlighted red.")
    ], widths=[2.5, 4.0])

    doc.add_heading("15. Save To SQL And CRUD Behavior", level=1)
    doc.add_paragraph("The save path uses CRUD-style operations: read, create/update, delete old rows for the same version, then insert current rows.")
    add_table(doc, ["Operation", "Where", "Meaning"], [
        ("Create/Update project", "GetOrCreateProject", "Inserts new SmaScheduleProjects row if needed, otherwise updates project metadata."),
        ("Delete old task rows", "SaveProject", "Deletes SmaScheduleTasks for the same ProjectId and VersionNumber only."),
        ("Delete old assignment rows", "SaveProject", "Deletes TaskAssignment rows for the same ProjectId and VersionNumber only."),
        ("Insert schedule task rows", "SaveProject", "Inserts one row per ScheduleTask into SmaScheduleTasks."),
        ("Insert assignment rows", "SaveTaskAssignments", "Expands task allocations into per employee/date rows in TaskAssignment."),
        ("Read saved schedules", "LoadSavedProjectSchedule / LoadSavedProjectScheduleByProjectCode", "Rebuilds a scheduler schedule from saved SQL rows.")
    ], widths=[1.9, 2.1, 2.5])
    add_code(doc, "DELETE FROM dbo.SmaScheduleTasks WHERE ProjectId = @ProjectId AND ISNULL(VersionNumber, '') = @VersionNumber\nDELETE FROM dbo.TaskAssignment WHERE ProjectId = @ProjectId AND ISNULL(VersionNumber, '') = @VersionNumber")
    doc.add_paragraph("This means tasks are not deleted only by Project ID. The same project can keep separate rows for different versions, as long as VersionNumber is correct.")

    doc.add_heading("16. SqlProjectRepository Function Map", level=1)
    add_table(doc, ["Function group", "Main functions", "Purpose"], [
        ("Connection/schema", "TestConnection, CreateConnection, EnsureSchema, TableExists, ColumnExists", "Open SQL connections and make sure expected local tables exist."),
        ("Project lookup", "GetProjectPlanningInfo, ListRecentScheduledProjects, LoadTemplateProjects", "Find project metadata and recent scheduled project lists."),
        ("Task templates", "LoadTaskCatalog, LoadTaskTemplates, LoadTaskTemplatesFromDatabaseDesign", "Load Task_Template rows for the scheduler."),
        ("Saved schedules", "LoadSavedProjectSchedule, LoadSavedProjectScheduleByProjectCode, LoadTasksByProjectId", "Rebuild schedules from SQL."),
        ("Saving", "SaveProject, GetOrCreateProject, SaveTaskAssignments, BuildTaskAssignmentRows", "Persist current schedule and capacity assignment rows."),
        ("Employees", "LoadEmployees, LoadEmployeeIdLookup, LoadEmployeeNameLookupFromMaster", "Map employee names to employee IDs and load active employees."),
        ("Capacity", "LoadCapacityPlanningData, LoadResourceAvailability, LoadAvailabilityFromEmployeeCapacity", "Load available and planned hours for date ranges."),
        ("Conversion helpers", "SqlBoolean, SqlText, SqlNullableDate, ProjectSizeNameFromSql, ProjectTypeFromVersion", "Convert SQL values into safe VB values.")
    ], widths=[1.5, 2.9, 2.1])

    doc.add_heading("17. Data Model Meaning", level=1)
    add_table(doc, ["Class", "Meaning"], [
        ("ScheduleTask", "One task row in the scheduler grid. It includes dates, duration, predecessor, assigned resources, allocation strings, total resource hours, module ID, and planner task ID."),
        ("TaskCatalogItem", "One SQL task template row. It stores the task name, dependency, hours for each project size, assignee, project type, report type, and task order."),
        ("LiveProjectItem", "One project selected from SQL. It carries project metadata from the planner into the scheduler."),
        ("ProjectLibraryItem", "One saved/recent project displayed in the planner library."),
        ("SqlProjectPlanningInfo", "Raw SQL planning details from Version_Table and Table_Project_Tracking."),
        ("SqlCapacityPlanningData", "Container for employees, availability rows, and planned assignment rows used by Capacity Planning.")
    ], widths=[1.8, 4.7])

    doc.add_heading("18. Theme And Project Flags", level=1)
    add_table(doc, ["Condition", "Effect"], [
        ("BRE/ROL/Within without pointcloud", "Uses one project-state theme."),
        ("Pointcloud true", "Uses pointcloud project-state theme."),
        ("Tech Pack or Deed Profile true", "Uses detail-heavy project-state theme."),
        ("Shadow Analysis true", "Uses shadow analysis project-state theme."),
        ("Urgent Small Projects true", "Uses urgent project-state theme."),
        ("Daily capacity overrun", "Capacity/risk theme can show warning colors.")
    ], widths=[2.5, 4.0])
    doc.add_paragraph("Main functions: ApplyProjectStateTheme, DetermineProjectStateTheme, ApplyTheme, ApplyGridTheme, StyleUsageSummaryCell, StylePendingHoursCell, StyleBlockedCell.")

    doc.add_heading("19. Files To Read First", level=1)
    add_table(doc, ["Priority", "File", "Reason"], [
        ("1", "Program.vb", "Small and easy; shows app entry point."),
        ("2", "LoginForm.vb", "Shows login rules and button/Enter behavior."),
        ("3", "SMAPlannerForm.vb", "Shows Project ID SQL validation and opening scheduler."),
        ("4", "TaskCatalogService.vb", "Shows tasks are SQL-only at runtime."),
        ("5", "TaskCatalogItem.vb", "Shows project-size hours mapping."),
        ("6", "ScheduleTask.vb", "Shows each scheduler grid row property."),
        ("7", "SMASchedulerForm.vb", "Main UI and button logic; read by sections."),
        ("8", "ScheduleEngine.vb", "Scheduling math, dependencies, and 8-hour capacity."),
        ("9", "SqlProjectRepository.vb", "SQL CRUD and data conversion logic.")
    ], widths=[0.8, 2.1, 3.6])

    doc.add_heading("20. Common Questions And Answers", level=1)
    add_table(doc, ["Question", "Answer"], [
        ("Are forms created in VB.NET?", "Yes. The project uses VB.NET WinForms. Designer files define visual controls, and .vb files contain code-behind."),
        ("Are tasks directly loaded from SQL?", "Yes for runtime. TaskCatalogService calls SQL repository; no default task list should appear to users."),
        ("Why not default Small and 1.0?", "Because SQL is the production source. Missing version or project size should stop scheduling and show a message."),
        ("Where does resource name come from?", "From SQL employee tables through EmployeeCatalogService and SqlProjectRepository.LoadEmployees."),
        ("What happens if selected user has no hours?", "Capacity views calculate available/planned hours and style over-capacity or unavailable cells with warning colors."),
        ("What does Employee Capacity do?", "It gives per employee/date available hours. Capacity Planning compares this with planned task hours."),
        ("Is snapshot an image?", "In this code, snip/snapshot behavior captures a panel image for sharing resource utilization."),
        ("What does .json mean?", "JSON is a structured data format often used for configuration or metadata."),
        ("What does .md mean?", "Markdown is a readable text document format used for notes, guides, and GitHub documentation.")
    ], widths=[2.2, 4.3])

    doc.add_heading("21. Debugging Checklist", level=1)
    add_bullets(doc, [
        "If scheduler does not open, check SQL connection string in App.config.",
        "If Project ID is rejected, check Version_Table.[Project ID at SMA].",
        "If project opens but no tasks appear, check Task_Template rows for project type, report type, active flag, and size hours.",
        "If project size is missing, check Table_Project_Tracking.[Project Size].",
        "If employees do not appear, check Employees_Master active rows and firstname values.",
        "If capacity is wrong, check Employee_Capacity for available hours and Project Schedule Table for planned hours.",
        "If save overwrites unexpected tasks, check VersionNumber stored with SmaScheduleTasks and TaskAssignment.",
        "If weekend dates appear unexpectedly, check the Include Saturdays option."
    ])

    doc.add_heading("22. Build Command", level=1)
    doc.add_paragraph("Use this command from the project folder to verify the solution compiles:")
    add_code(doc, 'dotnet build ".\\SMA Planning Engine.sln"')

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "SMA Planning Engine Code Study Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
